from datetime import datetime

from app.models.schemas import AnalysisResult, ComparisonRanking, JobResponse, JobStatus, ProviderResult, RankingEntry, ScoreBreakdown
from app.services.export import export_job_csv, export_job_excel, export_job_pdf, export_job_word
from app.services.export_report import TEST_CASE_COLUMNS, build_comparison_report


def _sample_job() -> JobResponse:
    results = [
        ProviderResult(
            solution_id="groq_whisper_groq_gemma",
            label="Groq Whisper + Groq Gemma 4 26B A4B",
            stt_provider="groq_whisper",
            llm_provider="openrouter_gemma",
            status="completed",
            transcript="Customer called about appointment delay.",
            analysis=AnalysisResult(
                sentiment="negative",
                summary="Customer reported appointment delay and requested callback.",
                resolution_status="partially_resolved",
                confidence=0.86,
                notes="Clear audio",
                key_issues=["appointment delay"],
            ),
            overall_score=0.82,
        ),
        ProviderResult(
            solution_id="sarvam_stt_sarvam_llm",
            label="Sarvam STT + Sarvam LLM",
            stt_provider="sarvam_stt",
            llm_provider="sarvam_llm",
            status="failed",
            error="Sarvam LLM JSON parse failed",
        ),
    ]
    ranking = ComparisonRanking(
        winner=RankingEntry(
            rank=1,
            solution_id="groq_whisper_groq_gemma",
            label="Groq Whisper + Groq Gemma 4 26B A4B",
            overall_score=0.82,
            score_breakdown=ScoreBreakdown(overall=0.82),
            recommendation_reason="Best overall score",
        ),
        recommendation_summary="Use Groq Whisper + Groq Gemma for fastest reliable analysis.",
    )
    return JobResponse(
        job_id="job-export-test",
        status=JobStatus.COMPLETED,
        created_at=datetime.utcnow(),
        completed_at=datetime.utcnow(),
        audio_filename="demo-call.wav",
        call_reference="DEMO-001",
        results=results,
        ranking=ranking,
        results_ready=True,
        aggregate_status="partial",
    )


def test_build_comparison_report_structure():
    report = build_comparison_report(_sample_job())
    assert report.header.project_name
    assert report.summary.total_test_cases == 4
    assert report.summary.passed_count == 1
    assert report.summary.failed_count == 3
    assert len(report.test_cases) == 4
    assert len(report.comparison_rows) == 4
    assert len(report.observations) == 4
    assert report.test_cases[0].test_id == "TC-001"
    passed_row = next(r for r in report.test_cases if r.status == "Pass")
    assert passed_row.pipeline == "Groq Whisper + Groq Gemma 4 26B A4B"
    assert report.final_comparison is not None


def test_build_comparison_report_includes_full_transcript():
    long_transcript = " ".join(["word"] * 400)
    job = _sample_job()
    job.results[0].transcript = long_transcript

    report = build_comparison_report(job)

    groq_row = next(
        r for r in report.test_cases
        if r.pipeline == "Groq Whisper + Groq Gemma 4 26B A4B"
    )
    assert groq_row.transcript == long_transcript
    assert "…" not in groq_row.transcript
    assert groq_row.summary == "Customer reported appointment delay and requested callback."


def test_export_job_csv_contains_sections():
    csv_text = export_job_csv(_sample_job())
    assert "Executive Summary" in csv_text
    assert "Test ID" in csv_text
    assert "TC-001" in csv_text
    assert "Final Recommendation" not in csv_text
    assert all(col in csv_text for col in TEST_CASE_COLUMNS)


def test_export_job_excel_is_xlsx_bytes():
    content = export_job_excel(_sample_job())
    assert content[:2] == b"PK"


def test_export_job_pdf_is_pdf_bytes():
    content = export_job_pdf(_sample_job())
    assert content.startswith(b"%PDF")


def test_export_job_word_is_docx_bytes():
    content = export_job_word(_sample_job())
    assert content[:2] == b"PK"


def test_export_job_word_is_call_analysis_only():
    """Word report focuses on call analysis — no ticket/CRM or pipeline comparison."""
    from io import BytesIO

    from docx import Document

    content = export_job_word(_sample_job())
    doc = Document(BytesIO(content))
    text = "\n".join(p.text for p in doc.paragraphs)

    assert "Call Analysis Report" in text
    assert "Customer reported appointment delay" in text
    assert "Sentiment:" in text
    assert "Confidence:" in text
    assert "Recommendation / Next Steps" in text or "Recommended" in text or "appointment delay" in text
    assert "Transcript" in text

    # Ticket / CRM / multi-pipeline comparison content must not appear.
    assert "Assigned team" not in text
    assert "Escalation status" not in text
    assert "Priority:" not in text
    assert "4-Solution Comparison" not in text
    assert "Job ID:" not in text
    assert "Final Comparison Result" not in text
    assert "Prepared by:" not in text


def test_export_job_word_includes_completed_analysis():
    content = export_job_word(_sample_job())
    assert len(content) > 1000
    from io import BytesIO
    from docx import Document

    doc = Document(BytesIO(content))
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "negative" in text.lower()
    assert "Customer called about appointment delay." in text

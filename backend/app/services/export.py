"""Client-ready export formats: CSV, Excel, PDF, Word."""
from __future__ import annotations

import csv
import io
import json

from openpyxl import Workbook
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

from app.models.schemas import JobResponse
from app.services.export_report import (
    COMPARISON_COLUMNS,
    TEST_CASE_COLUMNS,
    ComparisonReport,
    build_comparison_report,
)

# Brand colors aligned with application theme
COLOR_PRIMARY = colors.HexColor("#000000")
COLOR_SECONDARY = colors.HexColor("#e7000b")
COLOR_TEXT = colors.HexColor("#ffffff")
COLOR_MUTED = colors.HexColor("#666666")

EXCEL_HEADER_FILL = PatternFill("solid", fgColor="000000")
EXCEL_HEADER_FONT = Font(bold=True, color="FFFFFF", size=11)
EXCEL_PASS_FILL = PatternFill("solid", fgColor="E8F5E9")
EXCEL_FAIL_FILL = PatternFill("solid", fgColor="FFEBEE")
EXCEL_NEUTRAL_FILL = PatternFill("solid", fgColor="FFF8E1")
EXCEL_TITLE_FONT = Font(bold=True, size=14, color="000000")
EXCEL_SECTION_FONT = Font(bold=True, size=12, color="E7000B")


def export_job_json(job: JobResponse) -> str:
    payload = job.model_dump(mode="json")
    if payload.get("ranking"):
        payload["ranking"]["recommendation_summary"] = ""
        if payload["ranking"].get("winner"):
            payload["ranking"]["winner"]["recommendation_reason"] = ""
        payload["ranking"]["rankings"] = [
            {**entry, "recommendation_reason": ""}
            for entry in payload["ranking"].get("rankings", [])
        ]
    return json.dumps(payload, indent=2, default=str)


def export_job_csv(job: JobResponse) -> str:
    report = build_comparison_report(job)
    output = io.StringIO()
    writer = csv.writer(output, lineterminator="\n")

    h = report.header
    writer.writerow(["Project Name", h.project_name])
    writer.writerow(["Report Title", h.report_title])
    writer.writerow(["Date", h.date])
    writer.writerow(["Version", h.version])
    writer.writerow(["Prepared By", h.prepared_by])
    writer.writerow(["Job ID", h.job_id])
    writer.writerow(["Audio File", h.audio_filename])
    writer.writerow(["Call Reference", h.call_reference])
    writer.writerow([])

    s = report.summary
    writer.writerow(["Executive Summary"])
    writer.writerow(["Total Test Cases", s.total_test_cases])
    writer.writerow(["Passed", s.passed_count])
    writer.writerow(["Failed", s.failed_count])
    writer.writerow(["Neutral", s.neutral_count])
    writer.writerow(["Best-Performing Solution", s.best_solution])
    writer.writerow([])

    writer.writerow(TEST_CASE_COLUMNS)
    for row in report.test_cases:
        writer.writerow(row.as_list())

    return output.getvalue()


def _auto_width_sheet(ws, min_width: int = 12, max_width: int = 48) -> None:
    for col_cells in ws.columns:
        length = min_width
        col_letter = get_column_letter(col_cells[0].column)
        for cell in col_cells:
            if cell.value:
                length = max(length, min(len(str(cell.value)) + 2, max_width))
        ws.column_dimensions[col_letter].width = length


def _status_fill(status: str) -> PatternFill | None:
    if status == "Pass":
        return EXCEL_PASS_FILL
    if status == "Fail":
        return EXCEL_FAIL_FILL
    if status == "Neutral":
        return EXCEL_NEUTRAL_FILL
    return None


def export_job_excel(job: JobResponse) -> bytes:
    report = build_comparison_report(job)
    wb = Workbook()

    # Summary sheet
    ws_summary = wb.active
    ws_summary.title = "Executive Summary"
    h = report.header
    summary_rows = [
        ("Project Name", h.project_name),
        ("Report Title", h.report_title),
        ("Date", h.date),
        ("Version", h.version),
        ("Prepared By", h.prepared_by),
        ("Job ID", h.job_id),
        ("Audio File", h.audio_filename),
        ("Call Reference", h.call_reference),
        ("", ""),
        ("Total Test Cases", report.summary.total_test_cases),
        ("Passed", report.summary.passed_count),
        ("Failed", report.summary.failed_count),
        ("Neutral", report.summary.neutral_count),
        ("Best-Performing Solution", report.summary.best_solution),
    ]
    ws_summary["A1"] = h.report_title
    ws_summary["A1"].font = EXCEL_TITLE_FONT
    ws_summary.merge_cells("A1:B1")
    row_idx = 3
    for label, value in summary_rows:
        ws_summary.cell(row=row_idx, column=1, value=label).font = Font(bold=True)
        ws_summary.cell(row=row_idx, column=2, value=value)
        row_idx += 1
    _auto_width_sheet(ws_summary)

    # Test cases sheet
    ws_data = wb.create_sheet("Test Cases")
    ws_data.append(TEST_CASE_COLUMNS)
    for cell in ws_data[1]:
        cell.font = EXCEL_HEADER_FONT
        cell.fill = EXCEL_HEADER_FILL
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)

    status_col = TEST_CASE_COLUMNS.index("Status") + 1
    for row in report.test_cases:
        ws_data.append(row.as_list())
        status_cell = ws_data.cell(row=ws_data.max_row, column=status_col)
        fill = _status_fill(str(status_cell.value))
        if fill:
            for col in range(1, len(TEST_CASE_COLUMNS) + 1):
                ws_data.cell(row=ws_data.max_row, column=col).fill = fill

    ws_data.freeze_panes = "A2"
    ws_data.auto_filter.ref = ws_data.dimensions
    for row in ws_data.iter_rows(min_row=2):
        for cell in row:
            cell.alignment = Alignment(vertical="top", wrap_text=True)
    _auto_width_sheet(ws_data, min_width=14, max_width=52)

    buffer = io.BytesIO()
    wb.save(buffer)
    return buffer.getvalue()


def export_job_pdf(job: JobResponse) -> bytes:
    report = build_comparison_report(job)
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        leftMargin=0.6 * inch,
        rightMargin=0.6 * inch,
        topMargin=0.6 * inch,
        bottomMargin=0.6 * inch,
    )

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "ReportTitle",
        parent=styles["Heading1"],
        fontSize=18,
        textColor=COLOR_PRIMARY,
        spaceAfter=6,
    )
    subtitle_style = ParagraphStyle(
        "ReportSubtitle",
        parent=styles["Normal"],
        fontSize=11,
        textColor=COLOR_MUTED,
        spaceAfter=12,
    )
    section_style = ParagraphStyle(
        "Section",
        parent=styles["Heading2"],
        fontSize=13,
        textColor=COLOR_SECONDARY,
        spaceBefore=14,
        spaceAfter=8,
    )
    body_style = ParagraphStyle(
        "Body",
        parent=styles["Normal"],
        fontSize=10,
        leading=14,
    )

    story = []
    h = report.header
    story.append(Paragraph(h.project_name, title_style))
    story.append(Paragraph(h.report_title, subtitle_style))
    story.append(
        Paragraph(
            f"<b>Date:</b> {h.date} &nbsp;&nbsp; <b>Version:</b> {h.version} &nbsp;&nbsp; "
            f"<b>Prepared by:</b> {h.prepared_by}",
            body_style,
        )
    )
    story.append(
        Paragraph(
            f"<b>Job ID:</b> {h.job_id}<br/><b>Audio:</b> {h.audio_filename}<br/>"
            f"<b>Call Reference:</b> {h.call_reference}",
            body_style,
        )
    )
    story.append(Spacer(1, 0.2 * inch))

    s = report.summary
    story.append(Paragraph("Executive Summary", section_style))
    summary_data = [
        ["Metric", "Value"],
        ["Total Test Cases", str(s.total_test_cases)],
        ["Passed", str(s.passed_count)],
        ["Failed", str(s.failed_count)],
        ["Neutral", str(s.neutral_count)],
        ["Best-Performing Solution", s.best_solution],
    ]
    summary_table = Table(summary_data, colWidths=[2.2 * inch, 4.3 * inch])
    summary_table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), COLOR_PRIMARY),
                ("TEXTCOLOR", (0, 0), (-1, 0), COLOR_TEXT),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTSIZE", (0, 0), (-1, -1), 9),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f7f7f7")]),
            ]
        )
    )
    story.append(summary_table)
    story.append(Spacer(1, 0.25 * inch))

    story.append(Paragraph("Test Case Results", section_style))
    table_data = [TEST_CASE_COLUMNS]
    for row in report.test_cases:
        table_data.append(row.as_list())

    col_widths = [0.5 * inch, 0.75 * inch, 2.0 * inch, 0.6 * inch, 0.9 * inch, 1.2 * inch, 0.55 * inch, 0.45 * inch, 0.9 * inch]
    results_table = Table(table_data, colWidths=col_widths, repeatRows=1)
    table_style = TableStyle(
        [
            ("BACKGROUND", (0, 0), (-1, 0), COLOR_PRIMARY),
            ("TEXTCOLOR", (0, 0), (-1, 0), COLOR_TEXT),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE", (0, 0), (-1, -1), 7),
            ("GRID", (0, 0), (-1, -1), 0.25, colors.grey),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#fafafa")]),
        ]
    )
    status_idx = TEST_CASE_COLUMNS.index("Status")
    for i, row in enumerate(report.test_cases, start=1):
        if row.status == "Pass":
            table_style.add("BACKGROUND", (status_idx, i), (status_idx, i), colors.HexColor("#E8F5E9"))
        elif row.status == "Fail":
            table_style.add("BACKGROUND", (status_idx, i), (status_idx, i), colors.HexColor("#FFEBEE"))
        elif row.status == "Neutral":
            table_style.add("BACKGROUND", (status_idx, i), (status_idx, i), colors.HexColor("#FFF8E1"))
    results_table.setStyle(table_style)
    story.append(results_table)

    doc.build(story)
    return buffer.getvalue()


def _add_word_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value

    doc.add_paragraph()


def export_job_word(job: JobResponse) -> bytes:
    report = build_comparison_report(job)
    doc = Document()

    h = report.header
    title = doc.add_heading(h.project_name, level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT

    subtitle = doc.add_paragraph(h.report_title)
    subtitle.runs[0].font.size = Pt(14)
    subtitle.runs[0].font.color.rgb = RGBColor(102, 102, 102)

    doc.add_heading("Job Details", level=1)
    job_meta = doc.add_paragraph()
    job_meta.add_run(f"Date: {h.date}\n").bold = True
    job_meta.add_run(f"Version: {h.version}\n")
    job_meta.add_run(f"Prepared by: {h.prepared_by}\n")
    job_meta.add_run(f"Job ID: {h.job_id}\n")

    doc.add_heading("Audio Details", level=1)
    audio_meta = doc.add_paragraph()
    audio_meta.add_run(f"Audio file: {h.audio_filename}\n")
    audio_meta.add_run(f"Call reference: {h.call_reference}\n")

    doc.add_heading("Summary", level=1)
    s = report.summary
    summary_lines = [
        report.summary.comparison_note,
        f"Total pipelines evaluated: {s.total_test_cases}",
        f"Passed: {s.passed_count}",
        f"Failed: {s.failed_count}",
        f"Neutral: {s.neutral_count}",
        f"Best-performing solution: {s.best_solution}",
    ]
    for line in summary_lines:
        doc.add_paragraph(line, style="List Bullet")

    doc.add_heading("4-Solution Comparison", level=1)
    _add_word_table(
        doc,
        COMPARISON_COLUMNS,
        [row.as_list() for row in report.comparison_rows],
    )

    doc.add_heading("Per-Solution Observations", level=1)
    for obs in report.observations:
        doc.add_heading(obs.solution, level=2)
        lines = [
            f"Status: {obs.status}",
            f"Sentiment: {obs.sentiment}",
            f"Confidence: {obs.confidence}",
            f"Resolution: {obs.resolution_status}",
            f"Score: {obs.score}",
            f"Summary: {obs.summary}",
            f"Key issues: {obs.key_issues}",
            f"Recommended action: {obs.recommended_action}",
            f"Priority: {obs.action_priority}",
            f"Assigned team: {obs.assigned_team}",
            f"Escalation status: {obs.escalation_status}",
            f"Action items: {obs.action_items}",
            f"Notes: {obs.notes}",
        ]
        if obs.error and obs.error != "—":
            lines.append(f"Error: {obs.error}")
        for line in lines:
            doc.add_paragraph(line)
        doc.add_paragraph("Transcript:")
        doc.add_paragraph(obs.transcript)
        doc.add_paragraph()

    doc.add_heading("Final Comparison Result", level=1)
    if report.final_comparison:
        final = report.final_comparison
        doc.add_paragraph(
            f"Recommended solution: {final.winner} (score {final.winner_score})"
        )
        if final.rankings:
            doc.add_paragraph("Rankings:")
            for entry in final.rankings:
                doc.add_paragraph(entry, style="List Bullet")
    else:
        doc.add_paragraph("Ranking data was not available for this job.")

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
